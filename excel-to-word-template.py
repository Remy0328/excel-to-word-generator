"""
Excel 数据填充 Word 模板
========================
用法：
  python excel-to-word-template.py <excel文件> <word模板> <输出目录>
  
示例：
  python excel-to-word-template.py data.xlsx template.docx output/
  python excel-to-word-template.py                # 交互模式
"""

import os
import re
import sys
from pathlib import Path
import openpyxl
from docx import Document


def get_desktop():
    """获取桌面路径"""
    return Path.home() / "Desktop"


def select_file_interactive(title, filetypes):
    """交互式文件选择"""
    print(f"\n📂 {title}")
    print(f"   (直接按回车使用文件对话框，或手动输入路径)")
    
    # 尝试 tkinter
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        file_path = filedialog.askopenfilename(title=title, filetypes=filetypes)
        root.destroy()
        if file_path:
            return file_path
    except:
        pass
    
    # 降级：手动输入
    while True:
        path = input(f"   输入路径（或按回车打开对话框）: ").strip().strip('"')
        if not path:
            print("   ⚠️ 请输入文件路径")
            continue
        if os.path.exists(path):
            return path
        print(f"   ❌ 文件不存在: {path}")


def select_folder_interactive(title):
    """交互式文件夹选择"""
    print(f"\n📁 {title}")
    print(f"   (直接按回车使用文件夹对话框，或手动输入路径)")
    
    # 尝试 tkinter
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        folder = filedialog.askdirectory(title=title)
        root.destroy()
        if folder:
            return folder
    except:
        pass
    
    # 降级：手动输入
    while True:
        path = input(f"   输入路径（或按回车打开对话框）: ").strip().strip('"')
        if not path:
            print("   ⚠️ 请输入文件夹路径")
            continue
        if os.path.isdir(path):
            return path
        print(f"   ❌ 文件夹不存在: {path}")


def replace_placeholders(doc, data_dict):
    """替换 Word 文档中的所有占位符 {{字段名}}"""
    replaced_count = 0
    
    # 遍历所有段落
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if '{{' in run.text and '}}' in run.text:
                for key, value in data_dict.items():
                    placeholder = '{{' + key + '}}'
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
                        replaced_count += 1
    
    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if '{{' in run.text and '}}' in run.text:
                            for key, value in data_dict.items():
                                placeholder = '{{' + key + '}}'
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
                                    replaced_count += 1
    
    return replaced_count


def get_placeholder_fields(doc):
    """提取 Word 中所有占位符字段名"""
    fields = set()
    pattern = re.compile(r'\{\{(\w+)\}\}')
    
    for paragraph in doc.paragraphs:
        for match in pattern.finditer(paragraph.text):
            fields.add(match.group(1))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for match in pattern.finditer(paragraph.text):
                        fields.add(match.group(1))
    
    return sorted(fields)


def read_excel_headers(ws):
    """读取 Excel 第一行作为表头"""
    headers = []
    for cell in ws[1]:
        if cell.value is not None:
            headers.append(str(cell.value).strip())
    return headers


def read_excel_row(ws, row_idx):
    """读取 Excel 指定行数据"""
    data = {}
    headers = read_excel_headers(ws)
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=row_idx, column=col_idx)
        value = cell.value
        # 处理日期类型
        if hasattr(value, 'strftime'):
            value = value.strftime('%Y-%m-%d')
        data[header] = value if value is not None else ''
    return data


def generate_filename(data_dict, row_num):
    """生成输出文件名，优先使用 name/title 字段"""
    priority_fields = ['name', '姓名', 'title', '标题', 'filename', '文件名']
    
    for field in priority_fields:
        if field in data_dict and data_dict[field]:
            safe_name = re.sub(r'[\\/:*?"<>|]', '_', str(data_dict[field]))
            return f"{safe_name}.docx"
    
    return f"output_row{row_num}.docx"


def main():
    print("=" * 50)
    print("📄 Excel 数据 → Word 模板填充工具")
    print("=" * 50)
    
    # 解析命令行参数
    if len(sys.argv) >= 4:
        # 完整参数模式
        excel_path = sys.argv[1]
        template_path = sys.argv[2]
        output_folder = sys.argv[3]
        
        if not os.path.exists(excel_path):
            print(f"❌ Excel 文件不存在: {excel_path}")
            return
        if not os.path.exists(template_path):
            print(f"❌ Word 模板不存在: {template_path}")
            return
        if not os.path.isdir(output_folder):
            print(f"❌ 输出目录不存在: {output_folder}")
            return
            
    elif len(sys.argv) >= 2 and sys.argv[1] in ['-h', '--help', '/?']:
        print(__doc__)
        return
        
    else:
        # 交互模式
        excel_path = select_file_interactive(
            "选择 Excel 数据文件",
            [("Excel 文件", "*.xlsx *.xls"), ("所有文件", "*.*")]
        )
        print(f"   ✅ 已选择: {Path(excel_path).name}")
        
        template_path = select_file_interactive(
            "选择 Word 模板文件",
            [("Word 文档", "*.docx"), ("所有文件", "*.*")]
        )
        print(f"   ✅ 已选择: {Path(template_path).name}")
        
        output_folder = select_folder_interactive("选择输出文件夹")
        print(f"   ✅ 已选择: {output_folder}")
    
    # 加载 Excel 数据
    print("\n🔄 正在加载 Excel 数据...")
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
        ws = wb.active
        total_rows = ws.max_row - 1  # 减去表头行
        headers = read_excel_headers(ws)
        
        if not headers:
            print("❌ Excel 文件中没有找到表头行")
            return
        
        print(f"   ✅ 找到 {total_rows} 行数据，{len(headers)} 个字段")
        print(f"   字段: {', '.join(headers[:10])}", end='')
        if len(headers) > 10:
            print(f" ... 等 {len(headers)} 个")
        else:
            print()
    except Exception as e:
        print(f"❌ 读取 Excel 失败: {e}")
        return
    
    # 加载模板并提取占位符
    print("\n🔍 正在分析 Word 模板...")
    try:
        template_doc = Document(template_path)
        template_fields = get_placeholder_fields(template_doc)
        print(f"   ✅ 找到 {len(template_fields)} 个占位符: {template_fields}")
    except Exception as e:
        print(f"❌ 读取 Word 模板失败: {e}")
        return
    
    # 检查字段匹配
    matched_fields = [f for f in template_fields if f in headers]
    unmatched_in_template = [f for f in template_fields if f not in headers]
    unmatched_in_excel = [h for h in headers if h not in template_fields]
    
    if unmatched_in_template:
        print(f"\n⚠️ 模板中 {len(unmatched_in_template)} 个字段在 Excel 中没找到")
    
    if unmatched_in_excel:
        print(f"⚠️ Excel 中 {len(unmatched_in_excel)} 个字段在模板中未使用")
    
    if not matched_fields:
        print("\n❌ 模板占位符与 Excel 表头没有任何匹配！")
        print("   请确保 Excel 表头与模板占位符名称一致（如：{{姓名}}）")
        return
    
    print(f"\n✅ 字段匹配: {len(matched_fields)}/{len(template_fields)}")
    
    # 确认开始
    print(f"\n📋 确认:")
    print(f"   Excel: {Path(excel_path).name} ({total_rows} 行)")
    print(f"   模板: {Path(template_path).name}")
    print(f"   输出: {output_folder}")
    
    if total_rows > 100:
        confirm = input(f"\n将生成 {total_rows} 个文件，确认继续？(y/n): ").strip().lower()
        if confirm != 'y':
            print("❌ 已取消")
            return
    else:
        confirm = input("\n确认继续？(y/n): ").strip().lower()
        if confirm != 'y':
            print("❌ 已取消")
            return
    
    # 开始处理
    print(f"\n🚀 开始生成 Word 文件...")
    success_count = 0
    error_count = 0
    
    for row_num in range(2, ws.max_row + 1):
        try:
            # 读取 Excel 数据
            data = read_excel_row(ws, row_num)
            
            # 重新加载模板
            new_doc = Document(template_path)
            
            # 替换占位符
            replace_placeholders(new_doc, data)
            
            # 生成文件名
            filename = generate_filename(data, row_num)
            output_path = os.path.join(output_folder, filename)
            
            # 处理文件名冲突
            if os.path.exists(output_path):
                name, ext = os.path.splitext(filename)
                counter = 1
                while os.path.exists(output_path):
                    output_path = os.path.join(output_folder, f"{name}_{counter}{ext}")
                    counter += 1
            
            # 保存文件
            new_doc.save(output_path)
            success_count += 1
            print(f"   ✅ [{row_num-1}/{total_rows}] {filename}")
            
        except Exception as e:
            error_count += 1
            print(f"   ❌ 第 {row_num} 行处理失败: {e}")
    
    # 完成总结
    print("\n" + "=" * 50)
    if error_count > 0:
        print(f"🎉 完成！成功 {success_count} 个，失败 {error_count} 个")
    else:
        print(f"🎉 完成！成功生成 {success_count} 个 Word 文件")
    print(f"📁 输出目录: {output_folder}")
    print("=" * 50)


if __name__ == "__main__":
    main()
